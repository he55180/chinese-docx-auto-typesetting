# -*- coding: utf-8 -*-
"""
HSE公文自动化排版系统
原生python-docx模式，硬编码所有参数，符合GB/T 9704-2012标准
支持：章/节/级标题识别、中文编号矫正、无编号标题自动补全、标题正文混排拆分、
      编号层级重置、图片/表格保护
"""

__version__ = "2.0.3"

import re
import os
import sys
import argparse
import logging
import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, Mm, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from lxml import etree

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

# ===== 日志系统 =====
def setup_logging(log_file=None):
    log_format = '%(asctime)s - %(levelname)s - %(message)s'
    root_logger = logging.getLogger()
    root_logger.setLevel(logging.INFO)
    root_logger.handlers.clear()

    root_logger.addHandler(logging.StreamHandler(sys.stdout))
    if log_file:
        log_dir = os.path.dirname(log_file)
        if log_dir:
            os.makedirs(log_dir, exist_ok=True)
        fh = logging.FileHandler(log_file, 'w', encoding='utf-8')
        fh.setFormatter(logging.Formatter(log_format))
        root_logger.addHandler(fh)

    return root_logger

# ===== 硬编码参数（符合GB/T 9704-2012） =====
MARGIN_TOP = Mm(37)
MARGIN_BOTTOM = Mm(35)
MARGIN_LEFT = Mm(28)
MARGIN_RIGHT = Mm(26)

import platform as _platform
_IS_LINUX = _platform.system() == 'Linux'

FONT_MAIN_TITLE = 'WenQuanYi Zen Hei' if _IS_LINUX else '方正小标宋简体'
FONT_CHAPTER   = 'WenQuanYi Zen Hei' if _IS_LINUX else '方正小标宋简体'
FONT_SECTION   = 'Noto Sans CJK SC'   if _IS_LINUX else '黑体'
FONT_HEADING_1 = 'Noto Sans CJK SC'   if _IS_LINUX else '黑体'
FONT_HEADING_2 = 'Noto Sans CJK SC'   if _IS_LINUX else '楷体'
FONT_BODY      = 'Noto Sans CJK SC'   if _IS_LINUX else '仿宋_GB2312'

SIZE_MAIN_TITLE = Pt(22)
SIZE_CHAPTER = Pt(16)
SIZE_HEADING = Pt(16)
SIZE_BODY = Pt(16)

LINE_SPACING_FIXED = Pt(30)
FIRST_LINE_INDENT = Pt(32)
MAX_HEADING_LENGTH = 25

CHINESE_NUMERALS = ['〇', '一', '二', '三', '四', '五', '六', '七', '八', '九']

def to_chinese_numeral(n):
    if n <= 0 or n > 99:
        return ''
    if n < 10:
        return CHINESE_NUMERALS[n]
    tens = n // 10
    ones = n % 10
    if tens == 1:
        result = '十'
    else:
        result = CHINESE_NUMERALS[tens] + '十'
    if ones != 0:
        result += CHINESE_NUMERALS[ones]
    return result

# ===== 标题识别正则 =====
HEADING_PATTERNS = {
    'chapter': re.compile(r'^第[一二三四五六七八九十百]+章[：:\s]'),
    'section': re.compile(r'^第[一二三四五六七八九十百]+节[：:\s]'),
    'level1':  re.compile(r'^[一二三四五六七八九十]+[、．。）]'),
    'level1_shi': re.compile(r'^[一二三四五六七八九十]+是'),
    'level2':  re.compile(r'^（[一二三四五六七八九十]+）'),
    'level3':  re.compile(r'^\d+[.．、]'),
    'level4':  re.compile(r'^\(\d+\)'),   # v1.1.0 A1新增：(1) 格式四级标题
}

HEADING_STYLES = {
    'chapter':    {'font': FONT_CHAPTER,   'size': Pt(22), 'bold': False, 'align': WD_ALIGN_PARAGRAPH.CENTER},
    'section':    {'font': FONT_SECTION,   'size': Pt(16), 'bold': False, 'align': WD_ALIGN_PARAGRAPH.LEFT},
    'level1':     {'font': FONT_HEADING_1, 'size': Pt(16), 'bold': True,  'align': WD_ALIGN_PARAGRAPH.LEFT},
    'level1_shi': {'font': FONT_HEADING_1, 'size': Pt(18), 'bold': True,  'align': WD_ALIGN_PARAGRAPH.LEFT},
    'level2':     {'font': FONT_HEADING_2, 'size': Pt(16), 'bold': False, 'align': WD_ALIGN_PARAGRAPH.LEFT},
    'level3':     {'font': FONT_HEADING_1, 'size': Pt(18), 'bold': True,  'align': WD_ALIGN_PARAGRAPH.LEFT},
    'level2_auto':{'font': FONT_HEADING_2, 'size': Pt(16), 'bold': False, 'align': WD_ALIGN_PARAGRAPH.LEFT},
    'level4':     {'font': FONT_BODY,      'size': Pt(16), 'bold': False, 'align': WD_ALIGN_PARAGRAPH.LEFT},
    # v1.1.0 A1新增：level4 对应规范 3号仿宋体，内容分行排列
}

NEVER_BOLD_STYLES = {'FirstParagraph', '2', 'Normal', 'Body Text',
                     '正文', 'BodyText', 'Body'}

# ===== 标点符号修正模块（v2.0.0 新增） =====

# 全半角标点映射表
PUNCT_MAP = {
    # 引号
    '"': '\u201c',   # 左双引号
    '"': '\u201d',   # 右双引号
    "'": '\u2018',   # 左单引号
    "'": '\u2019',   # 右单引号
    # 括号
    '(': '\uff08',   # 全角左括号
    ')': '\uff09',   # 全角右括号
    # 标点
    ',': '\uff0c',   # 全角逗号
    ';': '\uff1b',   # 全角分号
    ':': '\uff1a',   # 全角冒号
    '!': '\uff01',   # 全角感叹号
    '?': '\uff1f',   # 全角问号
}

# 不替换的安全上下文：英文字母/数字前后的标点保留半角
import unicodedata

def _is_cjk(char):
    """判断是否为中文字符"""
    try:
        return unicodedata.name(char).startswith('CJK')
    except ValueError:
        return False

def fix_punctuation_in_text(text):
    """
    对单段文本执行标点修正。
    规则：仅当标点两侧至少一侧是中文字符或文本端点时，才替换为全角。
    英文句子内部的标点（两侧均为英文/数字）保留半角，避免误伤代码和英文。
    """
    if not text:
        return text
    result = list(text)
    length = len(text)
    for i, ch in enumerate(text):
        if ch not in PUNCT_MAP:
            continue
        left = text[i-1] if i > 0 else None
        right = text[i+1] if i < length - 1 else None
        left_is_en = left and (left.isascii() and (left.isalnum() or left in '._-'))
        right_is_en = right and (right.isascii() and (right.isalnum() or right in '._-'))
        # 两侧都是英文/数字：不替换（保护英文句子和代码）
        if left_is_en and right_is_en:
            continue
        result[i] = PUNCT_MAP[ch]
    return ''.join(result)

def fix_punctuation(doc):
    """
    遍历文档所有段落和表格单元格，对每个 run 的文本执行标点修正。
    跳过：纯英文段落、代码块（等宽字体段落）。
    """
    fixed_count = 0
    for para in doc.paragraphs:
        # 跳过纯英文段落（段落文本中中文字符占比低于 10%）
        full_text = para.text
        if not full_text.strip():
            continue
        cjk_ratio = sum(1 for c in full_text if _is_cjk(c)) / len(full_text)
        if cjk_ratio < 0.1:
            continue
        for run in para.runs:
            original = run.text
            fixed = fix_punctuation_in_text(original)
            if fixed != original:
                run.text = fixed
                fixed_count += 1
    # 同步处理表格单元格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        original = run.text
                        fixed = fix_punctuation_in_text(original)
                        if fixed != original:
                            run.text = fixed
                            fixed_count += 1
    return fixed_count

# ===== 格式深度清洗模块（v2.0.0 新增） =====

from docx.shared import RGBColor

# 允许保留的字体颜色（黑色系），其余一律清除
_ALLOWED_COLORS = {None, RGBColor(0, 0, 0), RGBColor(0x1F, 0x1F, 0x1F)}

def clean_dirty_format(doc):
    """
    深度清洗文档格式残留：
    1. 清除段落首尾多余空格
    2. 清除非黑色字体颜色（颜色标注等）
    3. 清除正文段落的非必要下划线、斜体
    4. 不处理表格内容（保留表格原始样式）
    """
    cleaned_count = 0
    BODY_STYLE_NAMES = {'Normal', 'Body Text', '正文', 'BodyText', 'Body', 'FirstParagraph', '2'}

    for para in doc.paragraphs:
        # 1. 清除首尾空格
        for run in para.runs:
            if run.text != run.text.strip():
                # 只清除纯空格的 run，避免破坏首行缩进等
                if run.text.strip() == '':
                    run.text = ''
                cleaned_count += 1

        # 2 & 3. 仅对正文段落清洗颜色/下划线/斜体
        style_name = para.style.name if para.style else ''
        is_body = style_name in BODY_STYLE_NAMES or get_paragraph_type(para.text.strip()) == 'body'
        if not is_body:
            continue

        for run in para.runs:
            # 清除非黑色字体颜色
            if run.font.color.rgb not in _ALLOWED_COLORS:
                run.font.color.rgb = RGBColor(0, 0, 0)
                cleaned_count += 1
            # 清除下划线（公文正文不应有下划线）
            if run.font.underline:
                run.font.underline = False
                cleaned_count += 1
            # 清除斜体（公文正文不应有斜体）
            if run.font.italic:
                run.font.italic = False
                cleaned_count += 1

    return cleaned_count

# ===== 主标题智能识别 =====
MAIN_TITLE_KEYWORDS = ['关于', '通知', '规定', '办法', '方案', '报告', '意见', '请示']

def is_main_title(para_text, is_first_para):
    if not is_first_para:
        return False
    text = para_text.strip()
    if not text:
        return False
    # 仅匹配红头文号模式：〔2026〕 或 "2026年12号"
    if re.search(r'〔\d{4}〕|\d{4}年\d+号', text):
        return False
    if len(text) > 50:
        return False
    if any(kw in text for kw in MAIN_TITLE_KEYWORDS):
        return True
    return False  # v1.1.0 A2修复：仅命中关键词才判定为主标题，防止称呼语等误判

# ===== 无编号标题智能识别 =====
MEETING_META_KEYWORDS = ['会议时间', '会议地点', '主持人', '记录人', '出席', '列席', '参会人员',
                         '会议日期', '会议主题', '发', '送', '抄送', '签发']

def is_likely_heading(text):
    stripped = text.strip()
    if not stripped:
        return False
    if len(stripped) > MAX_HEADING_LENGTH:
        return False
    # 排除会议元数据行
    if any(stripped.startswith(kw) for kw in MEETING_META_KEYWORDS):
        return False
    # 排除无编号章节名（前言、结束语、附录、后记）
    if stripped in ('前言', '结束语', '附录', '后记'):
        return False
    # 排除文号行：括号开头的年份编号（〔2025〕第4号 / （2025年第4次））
    if re.match(r'^[（(\[{〔]\d{4}', stripped):
        return False
    # 排除行内含日期的行
    if re.search(r'\d{4}年\d{1,2}月\d{1,2}日', stripped):
        return False
    if stripped[-1] in '。！？；':
        return False
    if stripped[-1] in '：:':
        return False
    if re.match(r'^\d+[.．、)]', stripped):
        return False
    if re.match(r'^[（(]\d+[）)]', stripped):
        return False
    if re.match(r'^(组长|副组长|成员|负责人|主任|经理|主管|领导|处长|科长|部长|院长|校长|书记|总[工程师]|汇报人|报告人|日期)[:\：]', stripped):
        return False
    if re.match(r'^\d{4}年\d{1,2}月\d{1,2}日', stripped):
        return False
    if re.search(r'\d{4}年\d{1,2}月\d{1,2}日', stripped):
        return False
    if ('，' in stripped or '；' in stripped) and len(stripped) > 12:
        return False

    # 排除包含"/"的机构联合名称（如"A机构/B机构"）
    if '/' in stripped or '／' in stripped:
        return False

    # 排除含有"项目""中心""公司""机构""组织""单位"
    # 且字数超过8字的机构名称行（主标题前的发文机关）
    ORG_SUFFIXES = ('项目', '中心', '公司', '机构', '组织',
                    '单位', '部门', '办公室', '委员会', '工程')
    if len(stripped) > 8 and any(stripped.endswith(s) for s in ORG_SUFFIXES):
        return False

    return True

# ===== 数字标题上下文判断 =====
BODY_STYLES = {'2', '正文', 'FirstParagraph', 'Normal', 'Body Text', 'Body', 'BodyText'}

def is_level3_heading(text, prev_style, next_style):
    """数字小标题判断：需同时满足以数字+点开头，且上下文不是纯正文环境"""
    if not re.match(r'^\d+[\.、．]\s', text):
        return False
    if len(text.strip()) > 40:          # v2.0.1 字数超40视为正文列表，不做标题
        return False
    if prev_style in BODY_STYLES and next_style in BODY_STYLES:
        return False
    return True

def is_organ_line(text, is_first_or_second_para):
    """
    识别主标题前的发文机关行。
    特征：文档前两段、字数<=40、
    以机构后缀结尾或含'/'联合机构分隔符。
    """
    if not is_first_or_second_para:
        return False
    stripped = text.strip()
    if len(stripped) > 40:
        return False
    ORG_SUFFIXES = ('项目', '中心', '公司', '机构', '组织',
                    '单位', '部门', '办公室', '委员会', '工程',
                    '集团', '局', '处', '队', '所')
    has_slash = '/' in stripped or '／' in stripped
    has_suffix = any(stripped.endswith(s) for s in ORG_SUFFIXES)
    return has_slash or has_suffix

# ===== 标题识别（按优先级从高到低） =====
def get_paragraph_type(text, is_first_or_second_para=False):
    stripped = text.strip()
    if not stripped:
        return 'empty'

    if is_organ_line(stripped, is_first_or_second_para):
        return 'organ'

    for ptype, pattern in HEADING_PATTERNS.items():
        if pattern.match(stripped):
            return ptype

    if is_likely_heading(stripped):
        return 'level2_auto'

    return 'body'

def strip_numbering(text):
    text = re.sub(r'^第[一二三四五六七八九十百]+[章节][：:\s]*', '', text.strip())
    text = re.sub(r'^[一二三四五六七八九十]+[、．。）]', '', text)
    text = re.sub(r'^[一二三四五六七八九十]+是', '', text)
    text = re.sub(r'^（[一二三四五六七八九十]+）', '', text)
    text = re.sub(r'^\d+[.．、]', '', text)
    return text.strip()

# ===== 统一截断规则：仅句子终止符（主体）/ 含冒号（compact） =====
FIRST_PUNCT = re.compile(r'[。．！!？?：:]')
SPLIT_PUNCT_HEADING = re.compile(r'[。．！!？?]')  # 标题行不用冒号截断，保护副标题

def split_at_first_punct(text, is_heading=False):
    """找到第一个终止符截断。标题行排除冒号以保护'项目名：副标题。'结构"""
    pat = SPLIT_PUNCT_HEADING if is_heading else FIRST_PUNCT
    m = pat.search(text)
    if m:
        return text[:m.end()], text[m.end():]
    return text, ''

def is_heading_with_body(text):
    """只有一、、（一）、1. 这三种格式开头的段落才需要截断"""
    return bool(re.match(
        r'^[一二三四五六七八九十]+[、．]\s*\S|'
        r'^（[一二三四五六七八九十]+）\s*\S|'
        r'^\d+[\.]\s*\S',
        text.strip()
    ))

def should_add_prefix(para_text, level):
    """已有标准公文编号前缀的段落不重复添加"""
    existing_prefixes = {
        'chapter': r'^第[一二三四五六七八九十百]+章',
        'section': r'^第[一二三四五六七八九十百]+节',
        'level1':  r'^[一二三四五六七八九十]+[、．。）]',
        'level1_shi': r'^[一二三四五六七八九十]+是',
        'level2':  r'^（[一二三四五六七八九十]+）',
        'level3':  r'^\d+[.．、]',
    }
    if level in existing_prefixes:
        return not bool(re.match(existing_prefixes[level], para_text))
    return True

def get_heading_prefix(para_text, level, counter):
    """生成标题编号前缀"""
    if level == 'chapter':
        return f"第{to_chinese_numeral(counter)}章  "
    elif level == 'section':
        return f"第{to_chinese_numeral(counter)}节  "
    elif level == 'level1':
        return f"{to_chinese_numeral(counter)}、"
    elif level == 'level1_shi':
        return f"{to_chinese_numeral(counter)}是"
    elif level == 'level2' or level == 'level2_auto':
        return f"（{to_chinese_numeral(counter)}）"
    return ''

# ===== 层级编号追踪 =====
class NumberingTracker:
    LEVEL_ORDER = ['chapter', 'section', 'level1', 'level1_shi', 'level2']

    def __init__(self):
        self.counters = {k: 0 for k in self.LEVEL_ORDER}

    def reset_below(self, level):
        idx = self.LEVEL_ORDER.index(level) if level in self.LEVEL_ORDER else -1
        for l in self.LEVEL_ORDER[idx+1:]:
            self.counters[l] = 0

    def next(self, level):
        if level in self.counters:
            self.reset_below(level)
            self.counters[level] += 1
            return self.counters[level]
        return 0

# ===== 文档渲染引擎 =====
class DocumentFormatter:
    def __init__(self, source_path, output_path, add_pagenum=True):
        self.source_path = source_path
        self.output_path = output_path
        self.add_pagenum = add_pagenum
        self.logger = logging.getLogger(__name__)
        self.numbering = NumberingTracker()

    def _apply_font_and_bold(self, run, font_name, size_pt=None, bold=False):
        run.bold = bold
        if size_pt:
            run.font.size = size_pt

        r = run._element
        rPr = r.get_or_add_rPr()

        ascii_font = "Times New Roman"
        if font_name in (FONT_HEADING_1, FONT_SECTION):
            ascii_font = "SimHei"
        elif font_name in (FONT_MAIN_TITLE, FONT_CHAPTER):
            ascii_font = font_name

        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), font_name)
        rFonts.set(qn('w:ascii'), ascii_font)
        rFonts.set(qn('w:hAnsi'), ascii_font)
        rFonts.set(qn('w:hint'), 'eastAsia')

        if size_pt:
            half_pts = str(int(size_pt.pt * 2))
            sz = rPr.get_or_add_sz()
            sz.set(qn('w:val'), half_pts)

            from docx.oxml import OxmlElement
            szCs = rPr.find(qn('w:szCs'))
            if szCs is None:
                szCs = OxmlElement('w:szCs')
                rPr.append(szCs)
            szCs.set(qn('w:val'), half_pts)

    def _format_compact_para(self, para, text):
        """Compact段落：统一规则——第一个标点截断，前加粗后正文"""
        bold_part, normal_part = split_at_first_punct(text)
        if not normal_part.strip():
            run = para.add_run(text)
            self._apply_font_and_bold(run, FONT_BODY, size_pt=Pt(16), bold=True)
            return 'compact', bold_part[:20]
        run1 = para.add_run(bold_part)
        self._apply_font_and_bold(run1, FONT_BODY, size_pt=Pt(16), bold=True)
        run2 = para.add_run(normal_part)
        self._apply_font_and_bold(run2, FONT_BODY, size_pt=Pt(16), bold=False)
        return 'compact', bold_part[:20]

    def _format_paragraph(self, para, p_text, p_type):
        has_image = any('drawing' in child.tag.lower() or 'pict' in child.tag.lower()
                      for run in para.runs for child in run._element)
        if has_image:
            return 'image', "图片段落保持原样"

        for run in para.runs:
            run.text = ''

        # 主标题
        if p_type == 'main_title':
            run = para.add_run(p_text)
            self._apply_font_and_bold(run, FONT_MAIN_TITLE, size_pt=SIZE_MAIN_TITLE, bold=True)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            para.paragraph_format.line_spacing = LINE_SPACING_FIXED
            para.paragraph_format.first_line_indent = Pt(0)
            return 'main_title', p_text[:30]

        # 发文机关 (organ) v2.0.3 新增 - 仿宋 14pt，左对齐，无缩进，不加编号
        if p_type == 'organ':
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            para.paragraph_format.line_spacing = LINE_SPACING_FIXED
            pPr = para._element.get_or_add_pPr()
            self._remove_numpr(pPr)
            self._set_indent(pPr, first_line=0)
            run = para.add_run(p_text)
            self._apply_font_and_bold(run, FONT_BODY, size_pt=Pt(14), bold=False)
            return 'organ', p_text[:30]

        # compact：降级的数字列表项，冒号前加粗，其余正文
        if p_type == 'compact':
            tag, log_msg = self._format_compact_para(para, p_text)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            para.paragraph_format.line_spacing = LINE_SPACING_FIXED
            para.paragraph_format.first_line_indent = FIRST_LINE_INDENT
            return tag, log_msg

        # body / compact：只有标题行才截断，普通正文不截断
        if p_type in ('body', 'compact'):
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            para.paragraph_format.line_spacing = LINE_SPACING_FIXED
            pPr = para._element.get_or_add_pPr()
            self._remove_numpr(pPr)
            if is_heading_with_body(p_text):
                tag, log_msg = self._format_compact_para(para, p_text)
                if tag == 'compact':
                    self._set_indent(pPr, first_line=640)
                    return tag, log_msg
            # 普通正文：整段不加粗，首行缩进
            self._set_indent(pPr, first_line=640)
            run = para.add_run(p_text)
            self._apply_font_and_bold(run, FONT_BODY, size_pt=SIZE_BODY, bold=False)
            return 'body', ''

        # 章/节/级标题
        style = HEADING_STYLES.get(p_type)
        if not style:
            run = para.add_run(p_text)
            self._apply_font_and_bold(run, FONT_BODY, size_pt=SIZE_BODY, bold=False)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.first_line_indent = FIRST_LINE_INDENT
            return 'body', ''

        counter_key = 'level2' if p_type == 'level2_auto' else p_type
        counter = self.numbering.next(counter_key)
        add_prefix = should_add_prefix(p_text, p_type)
        # 始终先剥离编号前缀，避免前缀标点干扰 split_at_first_punct
        clean_text = strip_numbering(p_text)

        # 统一截断规则：除章/节外，所有标题按第一个标点拆分（前加粗，后正文）
        split_types = {'level1', 'level1_shi', 'level2', 'level2_auto', 'level3', 'level4'}  # v1.1.0 A1新增level4
        if p_type in split_types:
            bold_part, normal_part = split_at_first_punct(clean_text, is_heading=True)
            if add_prefix:
                display_prefix = get_heading_prefix(p_text, p_type, counter)
            else:
                # 提取原始编号前缀（如"一、"、"（一）"、"1."），避免编号丢失
                prefix_map = {'level1': r'^[一二三四五六七八九十]+[、．。）]',
                              'level1_shi': r'^[一二三四五六七八九十]+是',
                              'level2': r'^（[一二三四五六七八九十]+）',
                              'level3': r'^\d+[\.．、]'}
                pm = re.match(prefix_map.get(p_type, ''), p_text)
                display_prefix = pm.group() if pm else ''
            if not normal_part.strip():
                display_text = f"{display_prefix}{bold_part}" if display_prefix else bold_part
                run = para.add_run(display_text)
                self._apply_font_and_bold(run, style['font'], size_pt=style['size'], bold=style['bold'])
            else:
                run1 = para.add_run(f"{display_prefix}{bold_part}")
                self._apply_font_and_bold(run1, style['font'], size_pt=style['size'], bold=style['bold'])
                run2 = para.add_run(normal_part)
                self._apply_font_and_bold(run2, FONT_BODY, size_pt=SIZE_BODY, bold=False)
            para.alignment = style['align']
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            para.paragraph_format.line_spacing = LINE_SPACING_FIXED
            para.paragraph_format.first_line_indent = Pt(0)
            display_text = f"{display_prefix}{bold_part}"
        else:
            # 章/节标题：不加前缀，直接渲染
            display_text = p_text
            run = para.add_run(display_text)
            self._apply_font_and_bold(run, style['font'], size_pt=style['size'], bold=style['bold'])
            para.alignment = style['align']
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            para.paragraph_format.line_spacing = LINE_SPACING_FIXED
            para.paragraph_format.first_line_indent = Pt(0)

        # 所有标题/compact段落：移除 numPr
        pPr = para._element.get_or_add_pPr()
        self._remove_numpr(pPr)

        # level3 悬挂缩进（数字突出，文字对齐）
        if p_type == 'level3':
            self._set_indent(pPr, first_line=0)
            ind = pPr.find(f'{{{W_NS}}}ind')
            if ind is not None:
                ind.set(f'{{{W_NS}}}left', '640')
                ind.set(f'{{{W_NS}}}hanging', '320')
        elif p_type in split_types:
            self._set_indent(pPr, first_line=640)

        return p_type, display_text[:30] if add_prefix else clean_text[:30]

    def _remove_numpr(self, pPr):
        """移除 Word 自动列表编号（numPr），避免自动加 1. 2. 3."""
        for numpr in pPr.findall(f'{{{W_NS}}}numPr'):
            pPr.remove(numpr)

    def _set_indent(self, pPr, first_line=640):
        """写入 w:ind，仅写 firstLine，left/hanging 由 Word 样式继承"""
        for old in pPr.findall(f'{{{W_NS}}}ind'):
            pPr.remove(old)
        ind = etree.SubElement(pPr, f'{{{W_NS}}}ind')
        ind.set(f'{{{W_NS}}}firstLine', str(first_line))

    def _format_tables(self, doc):
        for table in doc.tables:
            for row_idx, row in enumerate(table.rows):
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if row_idx == 0:
                                self._apply_font_and_bold(run, FONT_HEADING_1, size_pt=Pt(10), bold=True)
                            else:
                                self._apply_font_and_bold(run, FONT_BODY, size_pt=Pt(10), bold=False)
        self.logger.info("[OK] 表格格式化完成")

    def _add_page_numbers(self, doc):
        try:
            section = doc.sections[0]
            footer = section.footer
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            run = footer_para.add_run()
            from docx.oxml import OxmlElement

            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            run._r.append(fldChar1)

            run = footer_para.add_run()
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = "PAGE"
            run._r.append(instrText)

            run = footer_para.add_run()
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            run._r.append(fldChar2)

            for run in footer_para.runs:
                self._apply_font_and_bold(run, FONT_BODY, bold=False)
                run.font.size = Pt(10)

            self.logger.info("[OK] 页码已添加")
        except Exception as e:
            self.logger.error(f"[ERR] 页码添加失败: {e}")

    def format(self):
        results = {'success': 0, 'skipped': 0, 'failed': 0, 'errors': []}

        self.logger.info("="*70)
        self.logger.info("【HSE公文自动化排版 - 执行开始】")
        self.logger.info(f"源文件: {self.source_path}")
        self.logger.info(f"输出文件: {self.output_path}")
        self.logger.info("="*70)

        try:
            self.logger.info("\n【第一阶段】智能分析与预检...")
            doc = Document(self.source_path)
            para_count = len(doc.paragraphs)
            table_count = len(doc.tables)
            self.logger.info(f"[OK] 文档读取: {para_count}段落 + {table_count}表格")
        except Exception as e:
            self.logger.error(f"[ERR] 文档读取失败: {e}")
            return results

        section = doc.sections[0]
        section.top_margin = MARGIN_TOP
        section.bottom_margin = MARGIN_BOTTOM
        section.left_margin = MARGIN_LEFT
        section.right_margin = MARGIN_RIGHT
        section.page_width = Mm(210)
        section.page_height = Mm(297)

        self.logger.info("\n【第二阶段】服务格式渲染...")
        self.numbering = NumberingTracker()
        is_first_para = True

        TYPE_LABELS = {
            'chapter': '章标题', 'section': '节标题',
            'level1': '一级标题', 'level1_shi': '一是标题',
            'level2': '二级标题', 'level2_auto': '二级标题-自动',
            'level3': '三级标题', 'level4': '四级标题',  # v1.1.0 A1新增
            'main_title': '大标题', 'image': '跳过',
            'compact': '紧凑段落',
        }

        # 第一遍：收集段落信息（含样式名用于上下文判断）
        para_items = []
        is_first_para = True
        in_main_title = False
        non_empty_para_count = 0
        for i, para in enumerate(doc.paragraphs):
            p_text = para.text.strip()
            if not p_text:
                continue
            is_first_or_second_para = (non_empty_para_count < 2)
            p_type = get_paragraph_type(p_text, is_first_or_second_para)
            
            # 多行主标题识别逻辑
            if is_main_title(p_text, is_first_para):
                p_type = 'main_title'
                is_first_para = False
                in_main_title = True
            elif in_main_title:
                is_heading = p_type in HEADING_PATTERNS
                if (not is_heading) and (not p_text.endswith(('：', ':'))) and (len(p_text) <= 50) and (not re.search(r'〔\d{4}〕|\d{4}年\d+号', p_text)):
                    p_type = 'main_title'
                else:
                    in_main_title = False
            
            non_empty_para_count += 1
            para_style = para.style.name if para.style else ''
            para_items.append({'para': para, 'text': p_text, 'type': p_type, 'style': para_style, 'index': i})

        # 第二遍：上下文判断后格式化
        for idx, item in enumerate(para_items):
            para = item['para']
            p_text = item['text']
            p_type = item['type']

            # level3 上下文降级判断
            if p_type == 'level3':
                prev_style = para_items[idx-1]['style'] if idx > 0 else ''
                next_style = para_items[idx+1]['style'] if idx < len(para_items)-1 else ''
                if not is_level3_heading(p_text, prev_style, next_style):
                    p_type = 'compact'  # 降级为紧凑段落

            try:
                p_type, log_msg = self._format_paragraph(para, p_text, p_type)
                results['success'] += 1

                if p_type == 'image':
                    results['skipped'] += 1
                    results['success'] -= 1
                    self.logger.info(f"  [{TYPE_LABELS.get(p_type, p_type)}] {log_msg}")
                elif p_type in TYPE_LABELS:
                    self.logger.info(f"  [{TYPE_LABELS[p_type]}] {log_msg}")
                elif p_type != 'body':
                    self.logger.info(f"  [{p_type}] {log_msg}")

            except Exception as e:
                results['failed'] += 1
                results['errors'].append(f"第{idx+1}段('{item['text'][:15]}'): {str(e)[:80]}")  # v1.1.0 B1修复：i→idx
                self.logger.warning(f"  [跳过] 段落{idx+1}('{item['text'][:15]}')处理失败: {e}")

        self._format_tables(doc)

        if self.add_pagenum:
            try:
                self._add_page_numbers(doc)
            except Exception as e:
                self.logger.error(f"[ERR] 页码添加失败: {e}")

        self.logger.info(f"\n【处理结果】成功: {results['success']}, 跳过: {results['skipped']}, 失败: {results['failed']}")
        self.logger.info(f"【编号统计】章:{self.numbering.counters['chapter']} 节:{self.numbering.counters['section']} "
                        f"一级:{self.numbering.counters['level1']} 一是:{self.numbering.counters['level1_shi']} 二级:{self.numbering.counters['level2']}")
        if results['errors']:
            self.logger.warning(f"【错误详情】共{len(results['errors'])}个段落处理失败")

        try:
            self.logger.info("\n【保存文档】")
            doc.save(self.output_path)
            self.logger.info(f"[OK] 保存成功: {self.output_path}")
        except Exception as e:
            self.logger.error(f"[ERR] 保存失败: {e}")
            return results

        self.logger.info("\n" + "="*70)
        self.logger.info("【[OK] 排版完成】")
        self.logger.info("="*70)
        return results

# ===== 命令行入口 =====
def main():
    parser = argparse.ArgumentParser(description='HSE公文自动化排版系统')
    parser.add_argument('input', help='输入 .docx 文件路径，或 JSON 配置文件')
    parser.add_argument('-o', '--output', default=None, help='输出路径（默认：输入文件名_formatted.docx）')
    parser.add_argument('-c', '--config', default=None, help='JSON配置文件（可选，用于设置日志和页码）')
    parser.add_argument('--no-pagenum', action='store_true', help='不添加页码')
    parser.add_argument(
        '--mode',
        choices=['full', 'diagnose', 'punct'],
        default='full',
        help=(
            'full（默认）：标点修正+清洗+完整排版；'
            'diagnose：仅检测问题，输出诊断报告，不修改文件；'
            'punct：仅做标点修正，保留原有排版格式'
        )
    )
    parser.add_argument('--version', action='version', version=f'%(prog)s {__version__}')
    args = parser.parse_args()

    if args.input.endswith('.json'):
        config_file = args.input
        try:
            with open(config_file, 'r', encoding='utf-8') as f:
                config = json.load(f)
        except Exception as e:
            print(f"[ERR] 配置文件错误: {e}")
            sys.exit(1)
        log_path = config.get('log_file')
        source = config.get('source')
        output = config.get('output')
        add_pagenum = config.get('add_pagenum', True)
    else:
        source = args.input
        if not os.path.isfile(source):
            print(f"[ERR] 文件不存在: {source}")
            sys.exit(1)

        if args.output:
            output = args.output
        else:
            base = os.path.splitext(os.path.basename(source))[0]
            output_dir = os.path.dirname(source) or '.'
            output = os.path.join(output_dir, f"{base}_formatted.docx")

        add_pagenum = not args.no_pagenum

        log_path = None
        if args.config:
            try:
                with open(args.config, 'r', encoding='utf-8') as f:
                    extra_config = json.load(f)
                log_path = extra_config.get('log_file')
            except Exception as e:
                print(f"[WARN] 配置文件读取失败，忽略: {e}")

    setup_logging(log_path)
    logger = logging.getLogger(__name__)

    if not (source and output):
        logger.error("[ERR] 必须指定输入和输出")
        sys.exit(1)

    if not os.path.isfile(source):
        logger.error(f"[ERR] 源文件不存在: {source}")
        sys.exit(1)

    output_dir = os.path.dirname(output)
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)

    # ===== 处理模式路由（v2.0.0 新增） =====
    mode = getattr(args, 'mode', 'full')

    if mode == 'diagnose':
        # 诊断模式：只读文档，输出问题报告，不修改文件
        logger.info("="*60)
        logger.info("【诊断模式】仅检测问题，不修改文件")
        logger.info("="*60)
        doc = Document(source)
        issues = []
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            # 检测标点问题
            fixed = fix_punctuation_in_text(text)
            if fixed != text:
                issues.append(f"  第{i+1}段 [标点问题] {text[:30]}...")
            # 检测首行缩进缺失
            pf = para.paragraph_format
            if pf.first_line_indent is None or pf.first_line_indent == 0:
                ptype = get_paragraph_type(text)
                if ptype == 'body':
                    issues.append(f"  第{i+1}段 [缺首行缩进] {text[:30]}...")
        logger.info(f"\n共发现 {len(issues)} 个问题：")
        for iss in issues:
            logger.info(iss)
        logger.info("\n【诊断完成】文件未被修改。")
        return

    if mode == 'punct':
        # 仅标点模式：只做标点修正，保留所有原有排版
        logger.info("="*60)
        logger.info("【标点修正模式】仅修正标点，保留原排版")
        logger.info("="*60)
        doc = Document(source)
        count = fix_punctuation(doc)
        doc.save(output)
        logger.info(f"[OK] 标点修正完成，共修正 {count} 处，输出：{output}")
        return

    # full 模式：先做标点+清洗，再走完整排版流程
    if mode == 'full':
        logger.info("[v2.0] 执行标点修正与格式清洗...")
        doc_pre = Document(source)
        punct_count = fix_punctuation(doc_pre)
        clean_count = clean_dirty_format(doc_pre)
        # 保存为临时文件，再传给 DocumentFormatter
        import tempfile
        tmp_fd, tmp_path = tempfile.mkstemp(suffix='.docx')
        os.close(tmp_fd)
        doc_pre.save(tmp_path)
        logger.info(f"[OK] 标点修正 {punct_count} 处，格式清洗 {clean_count} 处")
        source = tmp_path  # 替换 source，后续流程使用清洗后的临时文件

    formatter = DocumentFormatter(source, output, add_pagenum=add_pagenum)
    results = formatter.format()

    # 清理 full 模式的临时文件
    if mode == 'full' and 'tmp_path' in locals():
        try:
            os.remove(tmp_path)
        except Exception:
            pass
    if results['failed'] > results['success']:
        sys.exit(1)

if __name__ == '__main__':
    main()
